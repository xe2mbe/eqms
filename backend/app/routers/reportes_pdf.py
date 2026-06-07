import io
import os
import ssl
import smtplib
from datetime import datetime
from typing import List, Optional

from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import text

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
)

from app.database import get_db
from app import models
from app.auth import get_current_user, require_admin
from app.routers.configuracion import _load_smtp

router = APIRouter()

LOGO_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'static', 'LogoFMRE.png')

# ─── Colores FMRE ─────────────────────────────────────────────────────────────
FMRE_BLUE     = colors.HexColor('#1A569E')
FMRE_BLUE_ALT = colors.HexColor('#f0f5ff')
# Teal como color base para RS — menos vistoso que el morado, misma familia por plataforma
RS_TEAL      = colors.HexColor('#0f766e')
RS_TEAL_ALT  = colors.HexColor('#f0fdfa')
RS_TEAL_DARK = colors.HexColor('#134e4a')

# Paleta de tonos para separar visualmente cada plataforma RS
RS_PALETTE = [
    (colors.HexColor('#0f766e'), colors.HexColor('#f0fdfa'), colors.HexColor('#134e4a')),  # teal
    (colors.HexColor('#0e7490'), colors.HexColor('#ecfeff'), colors.HexColor('#083344')),  # cyan
    (colors.HexColor('#047857'), colors.HexColor('#ecfdf5'), colors.HexColor('#064e3b')),  # emerald
    (colors.HexColor('#0369a1'), colors.HexColor('#f0f9ff'), colors.HexColor('#0c4a6e')),  # sky
]
COL_WHITE     = colors.white
COL_LGRAY     = colors.lightgrey

# ─── Schemas ─────────────────────────────────────────────────────────────────

class SeccionesConfig(BaseModel):
    # RF
    resumen_general: bool = True
    por_zona: bool = True
    por_sistema: bool = True
    top_estaciones: int = 10
    por_estado: bool = True
    primera_vez: bool = False
    detalle_rf: bool = False
    # RS
    resumen_plataformas: bool = True
    desglose_plataformas: bool = True
    top_estaciones_rs: int = 10
    por_zona_rs: bool = True
    metricas_detalle: bool = False
    detalle_rs: bool = False


class PlantillaCreate(BaseModel):
    nombre: str
    tipo: str = 'rf'
    evento_rf_id: Optional[int] = None
    evento_rs_id: Optional[int] = None
    eventos_rf_ids: List[int] = []
    eventos_rs_ids: List[int] = []
    secciones: SeccionesConfig = SeccionesConfig()
    destinatarios: List[str] = []
    asunto_email: Optional[str] = "Estadísticas {evento} – {fecha}"
    activa: bool = True
    # Asignación
    rol_asignado: Optional[str] = None
    usuario_id: Optional[int] = None


class ProgramacionUpdate(BaseModel):
    destinatarios: List[str] = []
    prog_hora: Optional[str] = None
    prog_dia_semana: Optional[int] = None
    prog_activo: bool = False


class PlantillaOut(BaseModel):
    id: int
    nombre: str
    tipo: str
    evento_rf_id: Optional[int] = None
    evento_rs_id: Optional[int] = None
    eventos_rf_ids: List[int] = []
    eventos_rs_ids: List[int] = []
    eventos_rf_tipos: List[str] = []
    eventos_rs_tipos: List[str] = []
    secciones: dict
    destinatarios: List[str]
    asunto_email: Optional[str]
    activa: bool
    rol_asignado: Optional[str] = None
    usuario_id: Optional[int] = None
    usuario_nombre: Optional[str] = None
    prog_hora: Optional[str] = None
    prog_recurrencia: Optional[str] = None
    prog_dia_semana: Optional[int] = None
    prog_activo: bool = False
    prog_ultima_ejecucion: Optional[datetime] = None
    created_at: Optional[datetime]

    class Config:
        from_attributes = True


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _ev_filter(ev_ids: List[int], alias: str = 'r') -> tuple:
    """Devuelve (fragmento_sql, params) para filtrar por múltiples eventos."""
    if not ev_ids:
        return "", {}
    phs = ', '.join(f':_ev{i}' for i in range(len(ev_ids)))
    return f"AND {alias}.evento_id IN ({phs})", {f'_ev{i}': v for i, v in enumerate(ev_ids)}


def _ev_filter_bare(ev_ids: List[int], col: str = 'evento_id') -> tuple:
    """Igual que _ev_filter pero sin alias de tabla."""
    if not ev_ids:
        return "", {}
    phs = ', '.join(f':_ev{i}' for i in range(len(ev_ids)))
    return f"AND {col} IN ({phs})", {f'_ev{i}': v for i, v in enumerate(ev_ids)}


def _to_out(p: models.ReportePlantilla, db: Session = None, user_config=None) -> PlantillaOut:
    uc = user_config
    ev_rf_ids: List[int] = list(p.eventos_rf_ids or [])
    ev_rs_ids: List[int] = list(p.eventos_rs_ids or [])

    # Nombres de los eventos para mostrar en UI
    ev_rf_tipos: List[str] = []
    ev_rs_tipos: List[str] = []
    if db and ev_rf_ids:
        rows = db.execute(text("SELECT id, tipo FROM eventos WHERE id = ANY(:ids)"),
                          {"ids": ev_rf_ids}).fetchall()
        id_tipo = {r[0]: r[1] for r in rows}
        ev_rf_tipos = [id_tipo[i] for i in ev_rf_ids if i in id_tipo]
    if db and ev_rs_ids:
        rows = db.execute(text("SELECT id, tipo FROM eventos WHERE id = ANY(:ids)"),
                          {"ids": ev_rs_ids}).fetchall()
        id_tipo = {r[0]: r[1] for r in rows}
        ev_rs_tipos = [id_tipo[i] for i in ev_rs_ids if i in id_tipo]

    return PlantillaOut(
        id=p.id,
        nombre=p.nombre,
        tipo=p.tipo or 'rf',
        evento_rf_id=p.evento_rf_id,
        evento_rs_id=p.evento_rs_id,
        eventos_rf_ids=ev_rf_ids,
        eventos_rs_ids=ev_rs_ids,
        eventos_rf_tipos=ev_rf_tipos,
        eventos_rs_tipos=ev_rs_tipos,
        secciones=p.secciones or {},
        destinatarios=(uc.destinatarios if uc else None) or p.destinatarios or [],
        asunto_email=p.asunto_email,
        activa=p.activa,
        rol_asignado=p.rol_asignado,
        usuario_id=p.usuario_id,
        usuario_nombre=p.usuario.full_name if p.usuario else None,
        prog_hora=uc.prog_hora if uc else p.prog_hora,
        prog_recurrencia=p.prog_recurrencia,
        prog_dia_semana=uc.prog_dia_semana if uc else p.prog_dia_semana,
        prog_activo=(uc.prog_activo if uc else p.prog_activo) or False,
        prog_ultima_ejecucion=uc.prog_ultima_ejecucion if uc else p.prog_ultima_ejecucion,
        created_at=p.created_at,
    )


# ─── CRUD plantillas ─────────────────────────────────────────────────────────

@router.get("/plantillas", response_model=List[PlantillaOut])
def list_plantillas(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Admin ve todas; operadores solo las asignadas a su rol o usuario."""
    q = db.query(models.ReportePlantilla)
    if current_user.role != 'admin':
        from sqlalchemy import or_
        q = q.filter(
            or_(
                models.ReportePlantilla.rol_asignado == None,
                models.ReportePlantilla.rol_asignado == current_user.role,
                models.ReportePlantilla.usuario_id == current_user.id,
            )
        )
    plantillas = q.order_by(models.ReportePlantilla.nombre).all()

    # Cargar configs del usuario actual en un dict para lookup O(1)
    pids = [p.id for p in plantillas]
    user_configs = {}
    if pids:
        ucs = db.query(models.ReportePlantillaUserConfig).filter(
            models.ReportePlantillaUserConfig.plantilla_id.in_(pids),
            models.ReportePlantillaUserConfig.usuario_id == current_user.id,
        ).all()
        user_configs = {uc.plantilla_id: uc for uc in ucs}

    return [_to_out(p, db, user_configs.get(p.id)) for p in plantillas]


@router.post("/plantillas", response_model=PlantillaOut, status_code=201)
def create_plantilla(body: PlantillaCreate, db: Session = Depends(get_db), _=Depends(require_admin)):
    ev_rf_ids = body.eventos_rf_ids or ([body.evento_rf_id] if body.evento_rf_id else [])
    ev_rs_ids = body.eventos_rs_ids or ([body.evento_rs_id] if body.evento_rs_id else [])
    p = models.ReportePlantilla(
        nombre=body.nombre,
        tipo=body.tipo,
        evento_rf_id=ev_rf_ids[0] if ev_rf_ids else None,
        evento_rs_id=ev_rs_ids[0] if ev_rs_ids else None,
        eventos_rf_ids=ev_rf_ids,
        eventos_rs_ids=ev_rs_ids,
        secciones=body.secciones.model_dump(),
        destinatarios=body.destinatarios,
        asunto_email=body.asunto_email,
        activa=body.activa,
        rol_asignado=body.rol_asignado,
        usuario_id=body.usuario_id,
    )
    db.add(p)
    db.commit()
    db.refresh(p)
    return _to_out(p, db)


@router.put("/plantillas/{pid}", response_model=PlantillaOut)
def update_plantilla(pid: int, body: PlantillaCreate, db: Session = Depends(get_db), _=Depends(require_admin)):
    p = db.get(models.ReportePlantilla, pid)
    if not p:
        raise HTTPException(404, "Plantilla no encontrada")
    ev_rf_ids = body.eventos_rf_ids or ([body.evento_rf_id] if body.evento_rf_id else [])
    ev_rs_ids = body.eventos_rs_ids or ([body.evento_rs_id] if body.evento_rs_id else [])
    p.nombre = body.nombre
    p.tipo = body.tipo
    p.evento_rf_id = ev_rf_ids[0] if ev_rf_ids else None
    p.evento_rs_id = ev_rs_ids[0] if ev_rs_ids else None
    p.eventos_rf_ids = ev_rf_ids
    p.eventos_rs_ids = ev_rs_ids
    p.secciones = body.secciones.model_dump()
    p.destinatarios = body.destinatarios
    p.asunto_email = body.asunto_email
    p.activa = body.activa
    p.rol_asignado = body.rol_asignado
    p.usuario_id = body.usuario_id
    db.commit()
    db.refresh(p)
    return _to_out(p, db)


@router.put("/plantillas/{pid}/programacion", response_model=PlantillaOut)
def update_programacion(pid: int, body: ProgramacionUpdate,
                        db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    p = db.get(models.ReportePlantilla, pid)
    if not p:
        raise HTTPException(404, "Plantilla no encontrada")

    uc = db.query(models.ReportePlantillaUserConfig).filter_by(
        plantilla_id=pid, usuario_id=current_user.id
    ).first()
    if uc is None:
        uc = models.ReportePlantillaUserConfig(
            plantilla_id=pid,
            usuario_id=current_user.id,
        )
        db.add(uc)

    uc.destinatarios = body.destinatarios
    uc.prog_hora = body.prog_hora
    uc.prog_dia_semana = body.prog_dia_semana
    uc.prog_activo = body.prog_activo
    db.commit()
    db.refresh(uc)
    return _to_out(p, db, uc)


@router.delete("/plantillas/{pid}", status_code=204)
def delete_plantilla(pid: int, db: Session = Depends(get_db), _=Depends(require_admin)):
    p = db.get(models.ReportePlantilla, pid)
    if not p:
        raise HTTPException(404, "Plantilla no encontrada")
    db.delete(p)
    db.commit()


# ─── Último evento capturado ──────────────────────────────────────────────────

def _detectar_cluster(rows) -> tuple:
    """Recibe filas (fecha, cnt) en orden DESC y devuelve (fi, ff, fechas_asc)."""
    cluster = []
    prev = None
    for fecha, cnt in rows:
        if prev is None or (prev - fecha).days <= 2:
            cluster.append({"fecha": fecha.isoformat(), "count": int(cnt)})
            prev = fecha
        else:
            break
    if not cluster:
        return None, None, []
    ff = cluster[0]["fecha"]
    fi = cluster[-1]["fecha"]
    return fi, ff, sorted(cluster, key=lambda x: x["fecha"])


def _cluster_rf(db: Session, ev_ids: List[int], before_date=None) -> Optional[dict]:
    evf, evp = _ev_filter(ev_ids, alias='r')
    evf_bare, _ = _ev_filter_bare(ev_ids)
    date_cap = "AND fecha_reporte::date < :before_date" if before_date else ""

    evento_nombre = None
    if ev_ids:
        rows_ev = db.execute(text("SELECT tipo FROM eventos WHERE id = ANY(:ids)"),
                             {"ids": ev_ids}).fetchall()
        evento_nombre = " + ".join(r[0] for r in rows_ev) if rows_ev else None

    ultima = db.execute(text(f"""
        SELECT MAX(fecha_reporte::date) FROM reportes WHERE 1=1
        {evf_bare} {date_cap}
    """), {**evp, "before_date": before_date}).scalar()
    if ultima is None:
        return None

    rows = db.execute(text(f"""
        SELECT r.fecha_reporte::date AS fecha, COUNT(*) AS cnt
        FROM reportes r
        WHERE r.fecha_reporte::date <= :ultima {evf}
        GROUP BY r.fecha_reporte::date ORDER BY fecha DESC
    """), {**evp, "ultima": ultima}).fetchall()
    fi, ff, _ = _detectar_cluster(rows)
    if not fi:
        return None

    sistemas = db.execute(text(f"""
        SELECT s.id, s.nombre, s.color
        FROM sistemas s
        WHERE s.id IN (
            SELECT DISTINCT r.sistema_id FROM reportes r
            WHERE r.sistema_id IS NOT NULL
              AND r.fecha_reporte::date BETWEEN :fi AND :ff
              {evf}
        )
        ORDER BY s.nombre
    """), {**evp, "fi": fi, "ff": ff}).fetchall()

    origenes = []
    for sid, snombre, scolor in sistemas:
        fd = db.execute(text(f"""
            SELECT fecha_reporte::date AS fecha, COUNT(*) AS cnt
            FROM reportes r
            WHERE r.sistema_id = :sid
              AND r.fecha_reporte::date BETWEEN :fi AND :ff
              {evf}
            GROUP BY fecha_reporte::date ORDER BY fecha
        """), {**evp, "sid": sid, "fi": fi, "ff": ff}).fetchall()
        fechas = [{"fecha": f.isoformat(), "count": int(c)} for f, c in fd]
        if fechas:
            origenes.append({
                "nombre": snombre,
                "color": scolor or "#1677ff",
                "total": sum(x["count"] for x in fechas),
                "fi": fechas[0]["fecha"],
                "ff": fechas[-1]["fecha"],
                "fechas": fechas,
            })

    sin_sistema = db.execute(text(f"""
        SELECT fecha_reporte::date AS fecha, COUNT(*) AS cnt
        FROM reportes r
        WHERE r.sistema_id IS NULL
          AND r.fecha_reporte::date BETWEEN :fi AND :ff
          {evf}
        GROUP BY fecha_reporte::date ORDER BY fecha
    """), {**evp, "fi": fi, "ff": ff}).fetchall()
    if sin_sistema:
        fechas = [{"fecha": f.isoformat(), "count": int(c)} for f, c in sin_sistema]
        origenes.append({
            "nombre": "Sin sistema",
            "color": "#8c8c8c",
            "total": sum(x["count"] for x in fechas),
            "fi": fechas[0]["fecha"],
            "ff": fechas[-1]["fecha"],
            "fechas": fechas,
        })

    return {"fi": fi, "ff": ff, "evento_nombre": evento_nombre, "origenes": origenes}


def _cluster_rs(db: Session, ev_ids: List[int], before_date=None) -> Optional[dict]:
    evf, evp = _ev_filter(ev_ids, alias='r')
    evf_bare, _ = _ev_filter_bare(ev_ids)
    date_cap = "AND fecha_reporte::date < :before_date" if before_date else ""

    evento_nombre = None
    if ev_ids:
        rows_ev = db.execute(text("SELECT tipo FROM eventos WHERE id = ANY(:ids)"),
                             {"ids": ev_ids}).fetchall()
        evento_nombre = " + ".join(r[0] for r in rows_ev) if rows_ev else None

    ultima = db.execute(text(f"""
        SELECT MAX(fecha_reporte::date) FROM reportes_rs WHERE 1=1
        {evf_bare} {date_cap}
    """), {**evp, "before_date": before_date}).scalar()
    if ultima is None:
        return None

    rows = db.execute(text(f"""
        SELECT r.fecha_reporte::date AS fecha, COUNT(*) AS cnt
        FROM reportes_rs r
        WHERE r.fecha_reporte::date <= :ultima {evf}
        GROUP BY r.fecha_reporte::date ORDER BY fecha DESC
    """), {**evp, "ultima": ultima}).fetchall()
    fi, ff, _ = _detectar_cluster(rows)
    if not fi:
        return None

    plataformas = db.execute(text(f"""
        SELECT pl.id, pl.nombre, pl.color
        FROM plataformas_rs pl
        WHERE pl.id IN (
            SELECT DISTINCT r.plataforma_id FROM reportes_rs r
            WHERE r.fecha_reporte::date BETWEEN :fi AND :ff
              {evf}
        )
        ORDER BY pl.nombre
    """), {**evp, "fi": fi, "ff": ff}).fetchall()

    origenes = []
    for plid, plnombre, plcolor in plataformas:
        fd = db.execute(text(f"""
            SELECT fecha_reporte::date AS fecha, COUNT(*) AS cnt
            FROM reportes_rs r
            WHERE r.plataforma_id = :plid
              AND r.fecha_reporte::date BETWEEN :fi AND :ff
              {evf}
            GROUP BY fecha_reporte::date ORDER BY fecha
        """), {**evp, "plid": plid, "fi": fi, "ff": ff}).fetchall()
        fechas = [{"fecha": f.isoformat(), "count": int(c)} for f, c in fd]
        if fechas:
            origenes.append({
                "nombre": plnombre,
                "color": plcolor or "#722ed1",
                "total": sum(x["count"] for x in fechas),
                "fi": fechas[0]["fecha"],
                "ff": fechas[-1]["fecha"],
                "fechas": fechas,
            })

    return {"fi": fi, "ff": ff, "evento_nombre": evento_nombre, "origenes": origenes}


@router.get("/plantillas/{pid}/ultimo-evento")
def ultimo_evento(pid: int, db: Session = Depends(get_db), _=Depends(get_current_user)):
    """Devuelve nombre del evento y desglose por sistema/plataforma del último clúster."""
    p = db.get(models.ReportePlantilla, pid)
    if not p:
        raise HTTPException(404, "Plantilla no encontrada")

    tipo = p.tipo or 'rf'
    ev_rf = list(p.eventos_rf_ids or [])
    ev_rs = list(p.eventos_rs_ids or [])
    rf_data = _cluster_rf(db, ev_rf) if tipo in ('rf', 'ambos') else None
    rs_data = _cluster_rs(db, ev_rs) if tipo in ('rs', 'ambos') else None

    fechas_fi = [d["fi"] for d in [rf_data, rs_data] if d and d["fi"]]
    fechas_ff = [d["ff"] for d in [rf_data, rs_data] if d and d["ff"]]

    return {
        "fi": min(fechas_fi) if fechas_fi else None,
        "ff": max(fechas_ff) if fechas_ff else None,
        "rf": rf_data,
        "rs": rs_data,
    }


# ─── Datos RF ─────────────────────────────────────────────────────────────────

def _gather_rf(db: Session, ev_ids: List[int], fi: datetime, ff: datetime) -> dict:
    evf, evp = _ev_filter(ev_ids, alias='r')
    evf_bare, evp_bare = _ev_filter_bare(ev_ids)
    base = {"fi": fi, "ff": ff}

    total = db.execute(text(f"""
        SELECT COUNT(*) FROM reportes
        WHERE fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
    """), {**base, **evp_bare}).scalar() or 0

    estaciones = db.execute(text(f"""
        SELECT COUNT(DISTINCT indicativo) FROM reportes
        WHERE fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
    """), {**base, **evp_bare}).scalar() or 0

    estados_cnt = db.execute(text(f"""
        SELECT COUNT(DISTINCT estado) FROM reportes
        WHERE estado IS NOT NULL
          AND fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
    """), {**base, **evp_bare}).scalar() or 0

    por_zona = db.execute(text(f"""
        SELECT z.codigo, z.nombre, COUNT(*) AS total, COUNT(DISTINCT r.indicativo) AS ests
        FROM reportes r JOIN zonas z ON z.id = r.zona_id
        WHERE r.zona_id IS NOT NULL
          AND r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
        GROUP BY z.codigo, z.nombre ORDER BY total DESC
    """), {**base, **evp}).fetchall()

    por_sistema = db.execute(text(f"""
        SELECT s.codigo, COUNT(*) AS total
        FROM reportes r JOIN sistemas s ON s.id = r.sistema_id
        WHERE r.sistema_id IS NOT NULL
          AND r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
        GROUP BY s.codigo ORDER BY total DESC
    """), {**base, **evp}).fetchall()

    top_ests = db.execute(text(f"""
        SELECT r.indicativo, COALESCE(rx.nombre_completo,''), COALESCE(r.estado,''), COUNT(*) AS total
        FROM reportes r
        LEFT JOIN radioexperimentadores rx ON rx.indicativo = r.indicativo
        WHERE r.fecha_reporte <= :ff {evf}
        GROUP BY r.indicativo, rx.nombre_completo, r.estado
        ORDER BY total DESC LIMIT 50
    """), {"ff": ff, **evp}).fetchall()

    por_estado = db.execute(text(f"""
        SELECT estado, COUNT(*) AS total FROM reportes
        WHERE estado IS NOT NULL
          AND fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
        GROUP BY estado ORDER BY total DESC LIMIT 32
    """), {**base, **evp_bare}).fetchall()

    primera_vez = db.execute(text(f"""
        SELECT r.indicativo, COALESCE(rx.nombre_completo,''), COALESCE(r.estado,'')
        FROM reportes r
        LEFT JOIN radioexperimentadores rx ON rx.indicativo = r.indicativo
        WHERE r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
          AND NOT EXISTS (
              SELECT 1 FROM reportes r2
              WHERE r2.indicativo = r.indicativo {evf.replace('r.', 'r2.')}
                AND r2.fecha_reporte < :fi
          )
        GROUP BY r.indicativo, rx.nombre_completo, r.estado
        ORDER BY r.indicativo
    """), {**base, **evp}).fetchall()

    detalle = db.execute(text(f"""
        SELECT r.indicativo, COALESCE(rx.nombre_completo,''), r.senal,
               COALESCE(r.estado,''), COALESCE(s.codigo,''), COALESCE(z.codigo,''),
               r.fecha_reporte
        FROM reportes r
        LEFT JOIN radioexperimentadores rx ON rx.indicativo = r.indicativo
        LEFT JOIN sistemas s ON s.id = r.sistema_id
        LEFT JOIN zonas z ON z.id = r.zona_id
        WHERE r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
        ORDER BY r.fecha_reporte, r.indicativo
    """), {**base, **evp}).fetchall()

    # Etiqueta de eventos para el encabezado del PDF
    if ev_ids:
        rows_ev = db.execute(text("SELECT tipo FROM eventos WHERE id = ANY(:ids) ORDER BY tipo"),
                             {"ids": ev_ids}).fetchall()
        eventos_label = " + ".join(r[0] for r in rows_ev) if rows_ev else "Todos los eventos"
    else:
        eventos_label = "Todos los eventos"

    return {
        "_eventos_label": eventos_label,
        "total": int(total),
        "estaciones": int(estaciones),
        "estados_cnt": int(estados_cnt),
        "por_zona": [{"zona": r[0], "nombre": r[1], "total": r[2], "ests": r[3]} for r in por_zona],
        "por_sistema": [{"sistema": r[0], "total": r[1]} for r in por_sistema],
        "top_ests": [{"ind": r[0], "nombre": r[1], "estado": r[2], "total": r[3]} for r in top_ests],
        "por_estado": [{"estado": r[0], "total": r[1]} for r in por_estado],
        "primera_vez": [{"ind": r[0], "nombre": r[1], "estado": r[2]} for r in primera_vez],
        "detalle": [{"ind": r[0], "nombre": r[1], "senal": r[2], "estado": r[3], "sistema": r[4], "zona": r[5], "fecha": r[6]} for r in detalle],
    }


# ─── Datos RS ─────────────────────────────────────────────────────────────────

def _gather_rs(db: Session, ev_ids: List[int], fi: datetime, ff: datetime) -> dict:
    evf, evp = _ev_filter(ev_ids, alias='r')
    evf_bare, evp_bare = _ev_filter_bare(ev_ids)
    base = {"fi": fi, "ff": ff}

    total = db.execute(text(f"""
        SELECT COUNT(*) FROM reportes_rs
        WHERE fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
    """), {**base, **evp_bare}).scalar() or 0

    estaciones = db.execute(text(f"""
        SELECT COUNT(DISTINCT indicativo) FROM reportes_rs
        WHERE fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
    """), {**base, **evp_bare}).scalar() or 0

    por_plataforma = db.execute(text(f"""
        SELECT pl.nombre, COUNT(r.id) AS cnt
        FROM reportes_rs r
        JOIN plataformas_rs pl ON pl.id = r.plataforma_id
        WHERE r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
        GROUP BY pl.nombre ORDER BY cnt DESC
    """), {**base, **evp}).fetchall()

    por_zona = db.execute(text(f"""
        SELECT z.codigo, z.nombre, COUNT(*) AS total, COUNT(DISTINCT r.indicativo) AS ests
        FROM reportes_rs r JOIN zonas z ON z.id = r.zona_id
        WHERE r.zona_id IS NOT NULL
          AND r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
        GROUP BY z.codigo, z.nombre ORDER BY total DESC
    """), {**base, **evp}).fetchall()

    top_ests = db.execute(text(f"""
        SELECT r.indicativo, COALESCE(rx.nombre_completo,''), COUNT(*) AS total
        FROM reportes_rs r
        LEFT JOIN radioexperimentadores rx ON rx.indicativo = r.indicativo
        WHERE r.fecha_reporte <= :ff {evf}
        GROUP BY r.indicativo, rx.nombre_completo
        ORDER BY total DESC LIMIT 50
    """), {"ff": ff, **evp}).fetchall()

    metricas_rows = db.execute(text("""
        SELECT pl.nombre, e.valores
        FROM estadisticas_rs e
        JOIN plataformas_rs pl ON pl.id = e.plataforma_id
        WHERE e.fecha_reporte >= :fi AND e.fecha_reporte <= :ff
    """), base).fetchall()

    metricas: dict = {}
    for row in metricas_rows:
        pl_name = row[0]
        vals = row[1] or {}
        if pl_name not in metricas:
            metricas[pl_name] = {}
        for k, v in vals.items():
            metricas[pl_name][k] = metricas[pl_name].get(k, 0) + (v or 0)

    por_estado_rs = db.execute(text(f"""
        SELECT estado, COUNT(*) AS total FROM reportes_rs
        WHERE estado IS NOT NULL AND estado <> ''
          AND fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
        GROUP BY estado ORDER BY total DESC LIMIT 32
    """), {**base, **evp_bare}).fetchall()

    detalle_rs = db.execute(text(f"""
        SELECT r.indicativo, COALESCE(rx.nombre_completo,''), pl.nombre,
               COALESCE(r.estado,''), COALESCE(z.codigo,''), r.fecha_reporte,
               COALESCE(r.url_publicacion,'')
        FROM reportes_rs r
        LEFT JOIN radioexperimentadores rx ON rx.indicativo = r.indicativo
        JOIN plataformas_rs pl ON pl.id = r.plataforma_id
        LEFT JOIN zonas z ON z.id = r.zona_id
        WHERE r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
        ORDER BY r.fecha_reporte, r.indicativo
    """), {**base, **evp}).fetchall()

    # Per-platform breakdown
    platforms_in_range = db.execute(text(f"""
        SELECT DISTINCT pl.id, pl.nombre
        FROM reportes_rs r
        JOIN plataformas_rs pl ON pl.id = r.plataforma_id
        WHERE r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
        ORDER BY pl.nombre
    """), {**base, **evp}).fetchall()

    por_plataforma_data: dict = {}
    for pl_id, pl_nombre in platforms_in_range:
        pp_base = {"fi": fi, "ff": ff, "pl": pl_id}

        pl_total = db.execute(text(f"""
            SELECT COUNT(*) FROM reportes_rs
            WHERE plataforma_id = :pl AND fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
        """), {**pp_base, **evp_bare}).scalar() or 0

        pl_estaciones = db.execute(text(f"""
            SELECT COUNT(DISTINCT indicativo) FROM reportes_rs
            WHERE plataforma_id = :pl AND fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
        """), {**pp_base, **evp_bare}).scalar() or 0

        pl_top_ests = db.execute(text(f"""
            SELECT r.indicativo, COALESCE(rx.nombre_completo,''), COUNT(*) AS total
            FROM reportes_rs r
            LEFT JOIN radioexperimentadores rx ON rx.indicativo = r.indicativo
            WHERE r.plataforma_id = :pl AND r.fecha_reporte <= :ff {evf}
            GROUP BY r.indicativo, rx.nombre_completo
            ORDER BY total DESC LIMIT 50
        """), {"ff": ff, "pl": pl_id, **evp}).fetchall()

        pl_por_zona = db.execute(text(f"""
            SELECT z.codigo, z.nombre, COUNT(*) AS total, COUNT(DISTINCT r.indicativo) AS ests
            FROM reportes_rs r JOIN zonas z ON z.id = r.zona_id
            WHERE r.plataforma_id = :pl AND r.zona_id IS NOT NULL
              AND r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
            GROUP BY z.codigo, z.nombre ORDER BY total DESC
        """), {**pp_base, **evp}).fetchall()

        pl_metricas_rows = db.execute(text("""
            SELECT e.valores FROM estadisticas_rs e
            WHERE e.plataforma_id = :pl
              AND e.fecha_reporte >= :fi AND e.fecha_reporte <= :ff
        """), pp_base).fetchall()
        pl_metricas: dict = {}
        for row in pl_metricas_rows:
            for k, v in (row[0] or {}).items():
                pl_metricas[k] = pl_metricas.get(k, 0) + (v or 0)

        pl_metricas_defs = db.execute(text("""
            SELECT slug, nombre FROM metricas_rs
            WHERE plataforma_id = :pl AND is_active = true
            ORDER BY orden, id
        """), {"pl": pl_id}).fetchall()

        pl_por_estado = db.execute(text(f"""
            SELECT estado, COUNT(*) AS total FROM reportes_rs
            WHERE plataforma_id = :pl AND estado IS NOT NULL AND estado <> ''
              AND fecha_reporte::date >= :fi::date AND fecha_reporte::date <= :ff::date {evf_bare}
            GROUP BY estado ORDER BY total DESC LIMIT 32
        """), {**pp_base, **evp_bare}).fetchall()

        pl_detalle = db.execute(text(f"""
            SELECT r.indicativo, COALESCE(rx.nombre_completo,''),
                   COALESCE(r.estado,''), COALESCE(z.codigo,''), r.fecha_reporte
            FROM reportes_rs r
            LEFT JOIN radioexperimentadores rx ON rx.indicativo = r.indicativo
            LEFT JOIN zonas z ON z.id = r.zona_id
            WHERE r.plataforma_id = :pl AND r.fecha_reporte::date >= :fi::date AND r.fecha_reporte::date <= :ff::date {evf}
            ORDER BY r.fecha_reporte, r.indicativo
        """), {**pp_base, **evp}).fetchall()

        por_plataforma_data[pl_nombre] = {
            "total":        int(pl_total),
            "estaciones":   int(pl_estaciones),
            "top_ests":     [{"ind": r[0], "nombre": r[1], "total": r[2]} for r in pl_top_ests],
            "por_zona":     [{"zona": r[0], "nombre": r[1], "total": r[2], "ests": r[3]} for r in pl_por_zona],
            "metricas":     pl_metricas,
            "metricas_defs": [{"slug": r[0], "nombre": r[1]} for r in pl_metricas_defs],
            "por_estado":   [{"estado": r[0], "total": r[1]} for r in pl_por_estado],
            "detalle":      [{"ind": r[0], "nombre": r[1], "estado": r[2], "zona": r[3], "fecha": r[4]} for r in pl_detalle],
        }

    metricas_defs_by_platform = {
        pl_nombre: pl_data["metricas_defs"]
        for pl_nombre, pl_data in por_plataforma_data.items()
    }

    if ev_ids:
        rows_ev = db.execute(text("SELECT tipo FROM eventos WHERE id = ANY(:ids) ORDER BY tipo"),
                             {"ids": ev_ids}).fetchall()
        eventos_label = " + ".join(r[0] for r in rows_ev) if rows_ev else "Todos los eventos"
    else:
        eventos_label = "Todos los eventos"

    return {
        "_eventos_label": eventos_label,
        "total_rs": int(total),
        "estaciones_rs": int(estaciones),
        "por_plataforma": [{"nombre": r[0], "cnt": r[1]} for r in por_plataforma],
        "por_zona_rs": [{"zona": r[0], "nombre": r[1], "total": r[2], "ests": r[3]} for r in por_zona],
        "top_ests_rs": [{"ind": r[0], "nombre": r[1], "total": r[2]} for r in top_ests],
        "metricas": metricas,
        "metricas_defs_by_platform": metricas_defs_by_platform,
        "por_estado_rs": [{"estado": r[0], "total": r[1]} for r in por_estado_rs],
        "detalle_rs": [{"ind": r[0], "nombre": r[1], "plataforma": r[2], "estado": r[3], "zona": r[4], "fecha": r[5], "url": r[6]} for r in detalle_rs],
        "por_plataforma_data": por_plataforma_data,
    }


# ─── PDF helpers ──────────────────────────────────────────────────────────────

def _tbl_style(bg: colors.Color = None, alt: colors.Color = None) -> TableStyle:
    bg  = bg  or FMRE_BLUE
    alt = alt or FMRE_BLUE_ALT
    return TableStyle([
        ('BACKGROUND',    (0, 0), (-1, 0), bg),
        ('TEXTCOLOR',     (0, 0), (-1, 0), COL_WHITE),
        ('FONTNAME',      (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE',      (0, 0), (-1, -1), 9),
        ('GRID',          (0, 0), (-1, -1), 0.4, COL_LGRAY),
        ('TOPPADDING',    (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING',   (0, 0), (-1, -1), 6),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 6),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [COL_WHITE, alt]),
    ])


def _tbl_style_detail(bg: colors.Color = None, alt: colors.Color = None) -> TableStyle:
    bg  = bg  or FMRE_BLUE
    alt = alt or FMRE_BLUE_ALT
    return TableStyle([
        ('BACKGROUND',    (0, 0), (-1, 0), bg),
        ('TEXTCOLOR',     (0, 0), (-1, 0), COL_WHITE),
        ('FONTNAME',      (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME',      (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE',      (0, 0), (-1, -1), 8),
        ('GRID',          (0, 0), (-1, -1), 0.3, COL_LGRAY),
        ('TOPPADDING',    (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING',   (0, 0), (-1, -1), 4),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 4),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [COL_WHITE, alt]),
    ])


def _banner(label: str, tipo_tag: str, bg: colors.Color, styles) -> Table:
    s_tag = ParagraphStyle('BN_TAG', parent=styles['Normal'],
                           textColor=COL_WHITE, fontSize=11, fontName='Helvetica-Bold',
                           alignment=TA_CENTER)
    s_lbl = ParagraphStyle('BN_LBL', parent=styles['Normal'],
                           textColor=COL_WHITE, fontSize=12, fontName='Helvetica-Bold',
                           alignment=TA_CENTER)
    dark = colors.HexColor('#0f3d6e') if tipo_tag == 'RF' else RS_TEAL_DARK
    t = Table([[Paragraph(tipo_tag, s_tag), Paragraph(label, s_lbl)]],
              colWidths=[1.8 * cm, 15.2 * cm])
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (0, 0), dark),
        ('BACKGROUND',    (1, 0), (1, 0), bg),
        ('TOPPADDING',    (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    return t


def _pl_banner(label: str, bg: colors.Color, dark: colors.Color, styles) -> Table:
    s_icon = ParagraphStyle('PLI', parent=styles['Normal'],
                            textColor=COL_WHITE, fontSize=10, fontName='Helvetica-Bold',
                            alignment=TA_CENTER)
    s_lbl  = ParagraphStyle('PLL', parent=styles['Normal'],
                            textColor=COL_WHITE, fontSize=11, fontName='Helvetica-Bold')
    t = Table([[Paragraph('>>', s_icon), Paragraph(f' {label}', s_lbl)]],
              colWidths=[1.2 * cm, 15.8 * cm])
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (0, 0), dark),
        ('BACKGROUND',    (1, 0), (1, 0), bg),
        ('TOPPADDING',    (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING',   (1, 0), (1, 0), 8),
    ]))
    return t


# ─── Generación PDF ───────────────────────────────────────────────────────────

def _build_pdf(p: models.ReportePlantilla, data: dict, fi: datetime, ff: datetime) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            rightMargin=1.5 * cm, leftMargin=1.5 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm,
                            title=p.nombre)

    gen_time = datetime.now().strftime('%d/%m/%Y %H:%M')

    def _footer_cb(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 7.5)
        canvas.setFillColor(colors.grey)
        y = 0.65 * cm
        canvas.drawString(doc.leftMargin, y,
                          f"Generado por QMS – FMRE  ·  {gen_time}")
        canvas.drawCentredString(doc.pagesize[0] / 2, y,
                                 f"Página {canvas.getPageNumber()}")
        if os.path.exists(LOGO_PATH):
            canvas.drawImage(LOGO_PATH,
                             doc.pagesize[0] - doc.rightMargin - 1.4 * cm,
                             0.25 * cm,
                             width=1.4 * cm, height=1.0 * cm,
                             preserveAspectRatio=True, anchor='sw', mask='auto')
        canvas.restoreState()

    styles = getSampleStyleSheet()
    s_title   = ParagraphStyle('T', parent=styles['Title'],
                               textColor=FMRE_BLUE, fontSize=20,
                               spaceAfter=4, alignment=TA_CENTER)
    s_sub     = ParagraphStyle('S', parent=styles['Normal'],
                               textColor=colors.grey, fontSize=11,
                               spaceAfter=6, alignment=TA_CENTER)
    s_section = ParagraphStyle('H', parent=styles['Heading2'],
                               textColor=FMRE_BLUE, fontSize=12,
                               spaceBefore=14, spaceAfter=6)
    s_sec_rs  = ParagraphStyle('HRS', parent=styles['Heading2'],
                               textColor=RS_TEAL, fontSize=12,
                               spaceBefore=14, spaceAfter=6)
    s_cell    = ParagraphStyle('CELL', fontName='Helvetica', fontSize=8, leading=10)
    s_footer  = ParagraphStyle('F', parent=styles['Normal'],
                               fontSize=8, textColor=colors.grey,
                               alignment=TA_CENTER)

    tipo = p.tipo or 'rf'
    sec  = p.secciones or {}
    ev_rf_ids = list(p.eventos_rf_ids or [])
    ev_rs_ids = list(p.eventos_rs_ids or [])
    evento_rf_nombre = data.get('rf', {}).get('_eventos_label') or "Todos los eventos"
    evento_rs_nombre = data.get('rs', {}).get('_eventos_label') or "Todos los eventos"
    evento_nombre = evento_rf_nombre if tipo != 'rs' else evento_rs_nombre
    fecha_str = fi.strftime('%d/%m/%Y')
    if fi.date() != ff.date():
        fecha_str += f" – {ff.strftime('%d/%m/%Y')}"

    story = []
    details_story = []   # secciones detalladas van al final del reporte

    # ── Encabezado ────────────────────────────────────────────────────────────
    story.append(Paragraph("Federación Mexicana de Radioexperimentadores A.C.", s_sub))
    story.append(Paragraph(p.nombre, s_title))
    story.append(Paragraph(f"{evento_nombre}  ·  {fecha_str}", s_sub))
    story.append(HRFlowable(width="100%", thickness=2, color=FMRE_BLUE, spaceAfter=12))

    # ── Secciones RF ─────────────────────────────────────────────────────────
    if tipo in ('rf', 'ambos'):
        rf = data.get('rf', {})

        if tipo == 'ambos':
            story.append(_banner("Reportes de Radio Frecuencia", "RF", FMRE_BLUE, styles))
            story.append(Spacer(1, 0.3 * cm))

        if sec.get('resumen_general', True):
            story.append(Paragraph("Resumen General", s_section))
            t = Table([
                ['Métrica', 'Valor'],
                ['Total de QSOs registrados', str(rf.get('total', 0))],
                ['Estaciones participantes',  str(rf.get('estaciones', 0))],
                ['Estados representados',     str(rf.get('estados_cnt', 0))],
            ], colWidths=[13 * cm, 4 * cm])
            t.setStyle(_tbl_style())
            t.setStyle(TableStyle([('ALIGN', (1, 0), (1, -1), 'CENTER')]))
            story.append(t)

        if sec.get('por_zona', True) and rf.get('por_zona'):
            story.append(Paragraph("Actividad por Zona", s_section))
            total_rf_qsos = rf.get('total', 0) or 1
            rows = [['Zona', 'Nombre', 'QSOs', 'Estaciones', '%']]
            for r in rf['por_zona']:
                pct = f"{r['total'] / total_rf_qsos * 100:.1f}%"
                rows.append([r['zona'], r['nombre'], str(r['total']), str(r['ests']), pct])
            t = Table(rows, colWidths=[3 * cm, 7 * cm, 2.5 * cm, 2.5 * cm, 2 * cm])
            t.setStyle(_tbl_style())
            t.setStyle(TableStyle([('ALIGN', (2, 0), (-1, -1), 'CENTER')]))
            story.append(t)

        if sec.get('por_sistema', True) and rf.get('por_sistema'):
            story.append(Paragraph("Actividad por Sistema", s_section))
            total_s = sum(r['total'] for r in rf['por_sistema']) or 1
            rows = [['Sistema', 'QSOs', '%']]
            for r in rf['por_sistema']:
                rows.append([r['sistema'], str(r['total']), f"{r['total'] / total_s * 100:.1f}%"])
            t = Table(rows, colWidths=[10 * cm, 3.5 * cm, 3.5 * cm])
            t.setStyle(_tbl_style())
            t.setStyle(TableStyle([('ALIGN', (1, 0), (-1, -1), 'CENTER')]))
            story.append(t)

        top_n = int(sec.get('top_estaciones', 10))
        if top_n > 0 and rf.get('top_ests'):
            story.append(Paragraph(f"Top {top_n} Estaciones — Acumulado hasta {ff.strftime('%d/%m/%Y')}", s_section))
            regular = [r for r in rf['top_ests'] if not r['ind'].upper().startswith('SWL')]
            swl_ests = [r for r in rf['top_ests'] if r['ind'].upper().startswith('SWL')]
            rows = [['#', 'Indicativo', 'Operador', 'Estado', 'QSOs']]
            for i, r in enumerate(regular[:top_n], 1):
                rows.append([str(i), r['ind'], r['nombre'], r['estado'], str(r['total'])])
            sep_idx = None
            if swl_ests:
                sep_idx = len(rows)
                rows.append(['', '', 'Escuchas  —  SWL', '', ''])
                for r in swl_ests:
                    rows.append(['—', r['ind'], r['nombre'], r['estado'], str(r['total'])])
            t = Table(rows, colWidths=[1 * cm, 3 * cm, 6.5 * cm, 4 * cm, 2.5 * cm])
            t.setStyle(_tbl_style())
            t.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                ('ALIGN', (4, 0), (4, -1), 'CENTER'),
            ]))
            if sep_idx is not None:
                t.setStyle(TableStyle([
                    ('BACKGROUND',  (0, sep_idx), (-1, sep_idx), colors.HexColor('#d6e4ff')),
                    ('TEXTCOLOR',   (0, sep_idx), (-1, sep_idx), FMRE_BLUE),
                    ('FONTNAME',    (0, sep_idx), (-1, sep_idx), 'Helvetica-BoldOblique'),
                    ('SPAN',        (0, sep_idx), (-1, sep_idx)),
                    ('ALIGN',       (0, sep_idx), (-1, sep_idx), 'CENTER'),
                    ('TOPPADDING',  (0, sep_idx), (-1, sep_idx), 4),
                    ('BOTTOMPADDING', (0, sep_idx), (-1, sep_idx), 4),
                ]))
            story.append(t)

        if sec.get('por_estado', True) and rf.get('por_estado'):
            story.append(Paragraph("Actividad por Estado", s_section))
            total_rf_qsos2 = rf.get('total', 0) or 1
            rows_est = rf['por_estado']
            half = (len(rows_est) + 1) // 2
            left, right = rows_est[:half], rows_est[half:]
            tbl = [['Estado', 'QSOs', '%', 'Estado', 'QSOs', '%']]
            for i in range(half):
                lr = left[i]
                rr = right[i] if i < len(right) else None
                lr_pct = f"{lr['total'] / total_rf_qsos2 * 100:.1f}%"
                rr_estado = rr['estado'] if rr else ''
                rr_total  = str(rr['total']) if rr else ''
                rr_pct    = f"{rr['total'] / total_rf_qsos2 * 100:.1f}%" if rr else ''
                tbl.append([lr['estado'], str(lr['total']), lr_pct,
                             rr_estado, rr_total, rr_pct])
            t = Table(tbl, colWidths=[4.5*cm, 2*cm, 1.5*cm, 4.5*cm, 2*cm, 2.5*cm])
            t.setStyle(_tbl_style())
            t.setStyle(TableStyle([
                ('ALIGN',     (1, 0), (2, -1), 'CENTER'),
                ('ALIGN',     (4, 0), (5, -1), 'CENTER'),
                ('LINEAFTER', (2, 0), (2, -1), 0.8, FMRE_BLUE),
            ]))
            story.append(t)

        if sec.get('primera_vez', False) and rf.get('primera_vez'):
            story.append(Paragraph(
                f"Nuevas Estaciones — {len(rf['primera_vez'])} primera aparición", s_section))
            rows = [['Indicativo', 'Operador', 'Estado']]
            for r in rf['primera_vez']:
                rows.append([r['ind'], r['nombre'], r['estado']])
            t = Table(rows, colWidths=[3.5 * cm, 9 * cm, 4.5 * cm])
            t.setStyle(_tbl_style())
            story.append(t)

        if sec.get('detalle_rf', False) and rf.get('detalle'):
            details_story.append(Paragraph(
                f"Reporte Detallado RF — {len(rf['detalle'])} QSOs", s_section))
            rows = [['#', 'Fecha', 'Indicativo', 'Operador', 'Señal', 'Estado', 'Sistema', 'Zona']]
            for i, r in enumerate(rf['detalle'], 1):
                rows.append([str(i), r['fecha'].strftime('%d/%m/%Y'), r['ind'],
                              Paragraph(r['nombre'] or '', s_cell),
                              str(r['senal']),
                              Paragraph(r['estado'] or '', s_cell),
                              r['sistema'], r['zona']])
            t = Table(rows, colWidths=[0.5*cm, 2.0*cm, 2.2*cm, 4.3*cm, 1.0*cm, 4.0*cm, 1.5*cm, 1.5*cm])
            t.setStyle(_tbl_style_detail())
            t.setStyle(TableStyle([
                ('ALIGN',  (0, 0), (0, -1), 'CENTER'),
                ('ALIGN',  (4, 0), (4, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            details_story.append(t)

    # ── Secciones RS ─────────────────────────────────────────────────────────
    if tipo in ('rs', 'ambos'):
        rs = data.get('rs', {})

        if tipo == 'ambos':
            story.append(Spacer(1, 0.5 * cm))
            story.append(_banner("Redes Sociales", "RS", RS_TEAL, styles))
            story.append(Spacer(1, 0.3 * cm))

        # ── Resumen global ────────────────────────────────────────────────
        if sec.get('resumen_plataformas', True):
            story.append(Paragraph("Resumen Redes Sociales", s_sec_rs))
            t = Table([
                ['Participación RS', 'Total'],
                ['Total reportes RS', str(rs.get('total_rs', 0))],
            ], colWidths=[13 * cm, 4 * cm])
            t.setStyle(_tbl_style(RS_TEAL, RS_TEAL_ALT))
            t.setStyle(TableStyle([('ALIGN', (1, 0), (1, -1), 'CENTER')]))
            story.append(t)

            if rs.get('por_plataforma'):
                story.append(Spacer(1, 0.3 * cm))
                total_rs_cnt = rs.get('total_rs', 0) or 1
                rows = [['Plataforma', 'Estaciones', '%']]
                for r in rs['por_plataforma']:
                    pl_d = rs.get('por_plataforma_data', {}).get(r['nombre'], {})
                    pct = f"{r['cnt'] / total_rs_cnt * 100:.1f}%"
                    rows.append([r['nombre'], str(pl_d.get('estaciones', '—')), pct])
                t = Table(rows, colWidths=[11 * cm, 3 * cm, 3 * cm])
                t.setStyle(_tbl_style(RS_TEAL, RS_TEAL_ALT))
                t.setStyle(TableStyle([('ALIGN', (1, 0), (-1, -1), 'CENTER')]))
                story.append(t)

        # ── Desglose por plataforma ───────────────────────────────────────
        if sec.get('desglose_plataformas', True):
            top_n_rs = int(sec.get('top_estaciones_rs', 10))
            for idx, (pl_nombre, pl_d) in enumerate(rs.get('por_plataforma_data', {}).items()):
                pl_bg, pl_alt, pl_dark = RS_PALETTE[idx % len(RS_PALETTE)]
                s_sec_pl = ParagraphStyle(f'HPL{idx}', parent=styles['Heading2'],
                                          textColor=pl_bg, fontSize=12,
                                          spaceBefore=14, spaceAfter=6)
                story.append(Spacer(1, 0.4 * cm))
                story.append(_pl_banner(pl_nombre, pl_bg, pl_dark, styles))
                story.append(Spacer(1, 0.2 * cm))

                # Participación de esta plataforma
                story.append(Paragraph(f"Participación — {pl_nombre}", s_sec_pl))
                t = Table([
                    ['Participación', 'Total'],
                    ['Total reportes', str(pl_d['total'])],
                ], colWidths=[13 * cm, 4 * cm])
                t.setStyle(_tbl_style(pl_bg, pl_alt))
                t.setStyle(TableStyle([('ALIGN', (1, 0), (1, -1), 'CENTER')]))
                story.append(t)

                # Métricas reales de esta plataforma (solo si tiene métricas configuradas)
                if pl_d.get('metricas_defs'):
                    story.append(Paragraph(f"Métricas — {pl_nombre}", s_sec_pl))
                    rows = [['Métrica', 'Total']]
                    for defn in pl_d['metricas_defs']:
                        val = pl_d.get('metricas', {}).get(defn['slug'], 0)
                        rows.append([defn['nombre'], str(int(val))])
                    t = Table(rows, colWidths=[13 * cm, 4 * cm])
                    t.setStyle(_tbl_style(pl_bg, pl_alt))
                    t.setStyle(TableStyle([('ALIGN', (1, 0), (1, -1), 'CENTER')]))
                    story.append(t)

                # Top estaciones de esta plataforma
                if top_n_rs > 0 and pl_d.get('top_ests'):
                    story.append(Paragraph(f"Top {top_n_rs} Estaciones {pl_nombre} — Acumulado hasta {ff.strftime('%d/%m/%Y')}", s_sec_pl))
                    rows = [['#', 'Indicativo', 'Operador', 'Reportes']]
                    for i, r in enumerate(pl_d['top_ests'][:top_n_rs], 1):
                        rows.append([str(i), r['ind'], r['nombre'], str(r['total'])])
                    t = Table(rows, colWidths=[1 * cm, 3 * cm, 9.5 * cm, 3.5 * cm])
                    t.setStyle(_tbl_style(pl_bg, pl_alt))
                    t.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                        ('ALIGN', (3, 0), (3, -1), 'CENTER'),
                    ]))
                    story.append(t)

                # Actividad por zona de esta plataforma
                if sec.get('por_zona_rs', True) and pl_d.get('por_zona'):
                    story.append(Paragraph(f"Actividad por Zona — {pl_nombre}", s_sec_pl))
                    total_pl_zona = pl_d['total'] or 1
                    rows = [['Zona', 'Nombre', 'Estaciones', '%']]
                    for r in pl_d['por_zona']:
                        pct = f"{r['total'] / total_pl_zona * 100:.1f}%"
                        rows.append([r['zona'], r['nombre'], str(r['ests']), pct])
                    t = Table(rows, colWidths=[3 * cm, 9 * cm, 3 * cm, 2 * cm])
                    t.setStyle(_tbl_style(pl_bg, pl_alt))
                    t.setStyle(TableStyle([('ALIGN', (2, 0), (-1, -1), 'CENTER')]))
                    story.append(t)

                # Actividad por estado de esta plataforma
                if pl_d.get('por_estado'):
                    story.append(Paragraph(f"Actividad por Estado — {pl_nombre}", s_sec_pl))
                    total_pl = pl_d['total'] or 1
                    rows_est = pl_d['por_estado']
                    half = (len(rows_est) + 1) // 2
                    left_e, right_e = rows_est[:half], rows_est[half:]
                    tbl = [['Estado', 'Rep.', '%', 'Estado', 'Rep.', '%']]
                    for i in range(half):
                        lr = left_e[i]
                        rr = right_e[i] if i < len(right_e) else None
                        tbl.append([
                            lr['estado'], str(lr['total']), f"{lr['total']/total_pl*100:.1f}%",
                            rr['estado'] if rr else '', str(rr['total']) if rr else '',
                            f"{rr['total']/total_pl*100:.1f}%" if rr else '',
                        ])
                    t = Table(tbl, colWidths=[4.5*cm, 1.8*cm, 1.5*cm, 4.5*cm, 1.8*cm, 2.9*cm])
                    t.setStyle(_tbl_style(pl_bg, pl_alt))
                    t.setStyle(TableStyle([
                        ('ALIGN',     (1, 0), (2, -1), 'CENTER'),
                        ('ALIGN',     (4, 0), (5, -1), 'CENTER'),
                        ('LINEAFTER', (2, 0), (2, -1), 0.8, pl_bg),
                    ]))
                    story.append(t)

                # Detalle de esta plataforma → va al final del reporte
                if sec.get('detalle_rs', False) and pl_d.get('detalle'):
                    details_story.append(Paragraph(
                        f"Reporte Detallado — {pl_nombre} — {len(pl_d['detalle'])} reportes",
                        s_sec_pl))
                    rows = [['#', 'Fecha', 'Indicativo', 'Operador', 'Estado', 'Zona']]
                    for i, r in enumerate(pl_d['detalle'], 1):
                        rows.append([str(i), r['fecha'].strftime('%d/%m/%Y'), r['ind'],
                                     Paragraph(r['nombre'] or '', s_cell),
                                     Paragraph(r['estado'] or '', s_cell),
                                     r['zona']])
                    t = Table(rows, colWidths=[0.5*cm, 2.0*cm, 2.5*cm, 5.0*cm, 4.5*cm, 2.5*cm])
                    t.setStyle(_tbl_style_detail(pl_bg, pl_alt))
                    t.setStyle(TableStyle([
                        ('ALIGN',  (0, 0), (0, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    details_story.append(t)

        else:
            # ── Modo global (sin desglose) ────────────────────────────────
            if sec.get('metricas_detalle', False) and rs.get('metricas'):
                for pl_name, vals in rs['metricas'].items():
                    defs = rs.get('metricas_defs_by_platform', {}).get(pl_name, [])
                    if not defs:
                        continue
                    story.append(Paragraph(f"Métricas — {pl_name}", s_sec_rs))
                    rows = [['Métrica', 'Total']]
                    for defn in defs:
                        val = vals.get(defn['slug'], 0)
                        rows.append([defn['nombre'], str(int(val))])
                    t = Table(rows, colWidths=[13 * cm, 4 * cm])
                    t.setStyle(_tbl_style(RS_TEAL, RS_TEAL_ALT))
                    t.setStyle(TableStyle([('ALIGN', (1, 0), (1, -1), 'CENTER')]))
                    story.append(t)

            if sec.get('por_zona_rs', True) and rs.get('por_zona_rs'):
                story.append(Paragraph("Actividad por Zona (RS)", s_sec_rs))
                rows = [['Zona', 'Nombre', 'Estaciones']]
                for r in rs['por_zona_rs']:
                    rows.append([r['zona'], r['nombre'], str(r['ests'])])
                t = Table(rows, colWidths=[3 * cm, 11 * cm, 3 * cm])
                t.setStyle(_tbl_style(RS_TEAL, RS_TEAL_ALT))
                t.setStyle(TableStyle([('ALIGN', (2, 0), (-1, -1), 'CENTER')]))
                story.append(t)

            if rs.get('por_estado_rs'):
                story.append(Paragraph("Actividad por Estado (RS)", s_sec_rs))
                total_rs2 = rs.get('total_rs', 0) or 1
                rows_est = rs['por_estado_rs']
                half = (len(rows_est) + 1) // 2
                left_e, right_e = rows_est[:half], rows_est[half:]
                tbl = [['Estado', 'Rep.', '%', 'Estado', 'Rep.', '%']]
                for i in range(half):
                    lr = left_e[i]
                    rr = right_e[i] if i < len(right_e) else None
                    tbl.append([
                        lr['estado'], str(lr['total']), f"{lr['total']/total_rs2*100:.1f}%",
                        rr['estado'] if rr else '', str(rr['total']) if rr else '',
                        f"{rr['total']/total_rs2*100:.1f}%" if rr else '',
                    ])
                t = Table(tbl, colWidths=[4.5*cm, 1.8*cm, 1.5*cm, 4.5*cm, 1.8*cm, 2.9*cm])
                t.setStyle(_tbl_style(RS_TEAL, RS_TEAL_ALT))
                t.setStyle(TableStyle([
                    ('ALIGN',     (1, 0), (2, -1), 'CENTER'),
                    ('ALIGN',     (4, 0), (5, -1), 'CENTER'),
                    ('LINEAFTER', (2, 0), (2, -1), 0.8, RS_TEAL),
                ]))
                story.append(t)

            top_n_rs = int(sec.get('top_estaciones_rs', 10))
            if top_n_rs > 0 and rs.get('top_ests_rs'):
                story.append(Paragraph(f"Top {top_n_rs} Estaciones RS — Acumulado hasta {ff.strftime('%d/%m/%Y')}", s_sec_rs))
                rows = [['#', 'Indicativo', 'Operador', 'Reportes']]
                for i, r in enumerate(rs['top_ests_rs'][:top_n_rs], 1):
                    rows.append([str(i), r['ind'], r['nombre'], str(r['total'])])
                t = Table(rows, colWidths=[1 * cm, 3 * cm, 9.5 * cm, 3.5 * cm])
                t.setStyle(_tbl_style(RS_TEAL, RS_TEAL_ALT))
                t.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                    ('ALIGN', (3, 0), (3, -1), 'CENTER'),
                ]))
                story.append(t)

            if sec.get('detalle_rs', False) and rs.get('detalle_rs'):
                details_story.append(Paragraph(
                    f"Reporte Detallado RS — {len(rs['detalle_rs'])} reportes", s_sec_rs))
                rows = [['#', 'Fecha', 'Indicativo', 'Operador', 'Plataforma', 'Estado', 'Zona']]
                for i, r in enumerate(rs['detalle_rs'], 1):
                    rows.append([str(i), r['fecha'].strftime('%d/%m/%Y'), r['ind'],
                                  Paragraph(r['nombre'] or '', s_cell),
                                  Paragraph(r['plataforma'] or '', s_cell),
                                  Paragraph(r['estado'] or '', s_cell),
                                  r['zona']])
                t = Table(rows, colWidths=[0.5*cm, 2.0*cm, 2.0*cm, 3.5*cm, 2.8*cm, 4.2*cm, 2.0*cm])
                t.setStyle(_tbl_style_detail(RS_TEAL, RS_TEAL_ALT))
                t.setStyle(TableStyle([
                    ('ALIGN',  (0, 0), (0, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                details_story.append(t)

    # ── Secciones detalladas al final ────────────────────────────────────────
    if details_story:
        story.append(Spacer(1, 0.4 * cm))
        story.append(HRFlowable(width="100%", thickness=1.5, color=FMRE_BLUE, spaceAfter=8))
        s_det_title = ParagraphStyle('DT', parent=styles['Heading1'],
                                     textColor=FMRE_BLUE, fontSize=14,
                                     spaceBefore=6, spaceAfter=10, alignment=TA_CENTER)
        story.append(Paragraph("Reportes Detallados", s_det_title))
        story.extend(details_story)

    doc.build(story, onFirstPage=_footer_cb, onLaterPages=_footer_cb)
    buf.seek(0)
    return buf.read()


# ─── Generación Word ─────────────────────────────────────────────────────────

def _add_word_table(doc: DocxDocument, headers: list, rows: list,
                    header_color: str = '1A569E'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = str(h)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        r, g, b = int(header_color[:2],16), int(header_color[2:4],16), int(header_color[4:],16)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_color.upper())
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        tr = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            c = tr.cells[ci]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(8)
    doc.add_paragraph()


def _build_word(p: models.ReportePlantilla, data: dict, fi: datetime, ff: datetime) -> bytes:
    doc = DocxDocument()

    # Estilos base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    tipo = p.tipo or 'rf'
    sec  = p.secciones or {}
    evento_nombre = data.get('rf', data.get('rs', {})).get('_eventos_label') or p.nombre
    fecha_str = fi.strftime('%d/%m/%Y')
    if fi.date() != ff.date():
        fecha_str += f" – {ff.strftime('%d/%m/%Y')}"

    title = doc.add_heading(p.nombre, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Federación Mexicana de Radioexperimentadores A.C.")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2 = doc.add_paragraph(f"{evento_nombre}  ·  {fecha_str}")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    BLUE, TEAL = '1A569E', '0F766E'

    if tipo in ('rf', 'ambos'):
        rf = data.get('rf', {})
        if tipo == 'ambos':
            doc.add_heading('Reportes de Radio Frecuencia (RF)', level=1)

        if sec.get('resumen_general', True):
            doc.add_heading('Resumen General', level=2)
            _add_word_table(doc, ['Métrica', 'Valor'], [
                ['Total de QSOs registrados', rf.get('total', 0)],
                ['Estaciones participantes',  rf.get('estaciones', 0)],
                ['Estados representados',     rf.get('estados_cnt', 0)],
            ], BLUE)

        if sec.get('por_zona', True) and rf.get('por_zona'):
            doc.add_heading('Actividad por Zona', level=2)
            total_q = rf.get('total', 0) or 1
            _add_word_table(doc, ['Zona', 'Nombre', 'QSOs', 'Estaciones', '%'],
                [[r['zona'], r['nombre'], r['total'], r['ests'],
                  f"{r['total']/total_q*100:.1f}%"] for r in rf['por_zona']], BLUE)

        if sec.get('por_sistema', True) and rf.get('por_sistema'):
            doc.add_heading('Actividad por Sistema', level=2)
            total_s = sum(r['total'] for r in rf['por_sistema']) or 1
            _add_word_table(doc, ['Sistema', 'QSOs', '%'],
                [[r['sistema'], r['total'], f"{r['total']/total_s*100:.1f}%"]
                 for r in rf['por_sistema']], BLUE)

        top_n = int(sec.get('top_estaciones', 10))
        if top_n > 0 and rf.get('top_ests'):
            doc.add_heading(f'Top {top_n} Estaciones RF', level=2)
            _add_word_table(doc, ['#', 'Indicativo', 'Operador', 'Estado', 'QSOs'],
                [[i+1, r['ind'], r['nombre'], r['estado'], r['total']]
                 for i, r in enumerate(rf['top_ests'][:top_n])], BLUE)

        if sec.get('por_estado', True) and rf.get('por_estado'):
            doc.add_heading('Actividad por Estado', level=2)
            total_q2 = rf.get('total', 0) or 1
            _add_word_table(doc, ['Estado', 'QSOs', '%'],
                [[r['estado'], r['total'], f"{r['total']/total_q2*100:.1f}%"]
                 for r in rf['por_estado']], BLUE)

        if sec.get('primera_vez', False) and rf.get('primera_vez'):
            doc.add_heading('Nuevas Estaciones', level=2)
            _add_word_table(doc, ['Indicativo', 'Operador', 'Estado'],
                [[r['ind'], r['nombre'], r['estado']] for r in rf['primera_vez']], BLUE)

        if sec.get('detalle_rf', False) and rf.get('detalle'):
            doc.add_heading('Reporte Detallado RF', level=2)
            _add_word_table(doc, ['#', 'Fecha', 'Indicativo', 'Operador', 'Señal', 'Estado', 'Sistema', 'Zona'],
                [[i+1, r['fecha'].strftime('%d/%m/%Y'), r['ind'], r['nombre'],
                  r['senal'], r['estado'], r['sistema'], r['zona']]
                 for i, r in enumerate(rf['detalle'])], BLUE)

    if tipo in ('rs', 'ambos'):
        rs = data.get('rs', {})
        if tipo == 'ambos':
            doc.add_heading('Redes Sociales (RS)', level=1)

        if sec.get('resumen_plataformas', True):
            doc.add_heading('Resumen Redes Sociales', level=2)
            _add_word_table(doc, ['Métrica', 'Total'], [
                ['Total reportes RS', rs.get('total_rs', 0)],
                ['Estaciones',        rs.get('estaciones_rs', 0)],
            ], TEAL)
            if rs.get('por_plataforma'):
                total_rs = rs.get('total_rs', 0) or 1
                _add_word_table(doc, ['Plataforma', 'Reportes', '%'],
                    [[r['nombre'], r['cnt'], f"{r['cnt']/total_rs*100:.1f}%"]
                     for r in rs['por_plataforma']], TEAL)

        if sec.get('desglose_plataformas', True):
            top_n_rs = int(sec.get('top_estaciones_rs', 10))
            for pl_nombre, pl_d in rs.get('por_plataforma_data', {}).items():
                doc.add_heading(f'Plataforma: {pl_nombre}', level=2)
                _add_word_table(doc, ['Métrica', 'Total'],
                    [['Reportes', pl_d['total']], ['Estaciones', pl_d['estaciones']]], TEAL)
                if pl_d.get('metricas_defs'):
                    doc.add_heading(f'Métricas — {pl_nombre}', level=3)
                    _add_word_table(doc, ['Métrica', 'Total'],
                        [[d['nombre'], int(pl_d.get('metricas', {}).get(d['slug'], 0))]
                         for d in pl_d['metricas_defs']], TEAL)
                if top_n_rs > 0 and pl_d.get('top_ests'):
                    doc.add_heading(f'Top {top_n_rs} Estaciones — {pl_nombre}', level=3)
                    _add_word_table(doc, ['#', 'Indicativo', 'Operador', 'Reportes'],
                        [[i+1, r['ind'], r['nombre'], r['total']]
                         for i, r in enumerate(pl_d['top_ests'][:top_n_rs])], TEAL)
                if sec.get('por_zona_rs', True) and pl_d.get('por_zona'):
                    doc.add_heading(f'Actividad por Zona — {pl_nombre}', level=3)
                    total_pl = pl_d['total'] or 1
                    _add_word_table(doc, ['Zona', 'Nombre', 'Estaciones', '%'],
                        [[r['zona'], r['nombre'], r['ests'], f"{r['total']/total_pl*100:.1f}%"]
                         for r in pl_d['por_zona']], TEAL)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Helper email (reutilizado por enviar + scheduler) ────────────────────────

def _send_email_report(p: models.ReportePlantilla, pdf: bytes, data: dict,
                       fi: datetime, ff: datetime, db: Session, destinatarios_override=None):
    from app.routers.configuracion import _load_smtp
    destinatarios = destinatarios_override or p.destinatarios or []
    if not destinatarios:
        raise ValueError("Sin destinatarios")
    smtp = _load_smtp(db)
    if not smtp.host or not smtp.usuario or not smtp.password:
        raise ValueError("SMTP no configurado")

    evento    = data.get('rf', data.get('rs', {})).get('_eventos_label') or p.nombre
    fname     = f"qms_{evento.replace(' ', '_').lower()}_{fi.strftime('%Y%m%d')}.pdf"
    fecha_str = fi.strftime('%d/%m/%Y')
    asunto    = (p.asunto_email or "Estadísticas {evento} – {fecha}") \
        .replace("{evento}", evento).replace("{fecha}", fecha_str)

    rf_block = ''
    if 'rf' in data:
        rf = data['rf']
        rf_block = f"""<p><strong>Estadísticas RF:</strong></p><ul>
  <li>Total QSOs: <strong>{rf['total']}</strong></li>
  <li>Estaciones: <strong>{rf['estaciones']}</strong></li>
  <li>Estados representados: <strong>{rf['estados_cnt']}</strong></li>
</ul>"""
    rs_block = ''
    if 'rs' in data:
        rs = data['rs']
        plats = ', '.join(f"{r['nombre']} ({r['cnt']})" for r in rs.get('por_plataforma', []))
        rs_block = f"""<p><strong>Estadísticas RS:</strong></p><ul>
  <li>Total reportes RS: <strong>{rs['total_rs']}</strong></li>
  <li>Estaciones: <strong>{rs['estaciones_rs']}</strong></li>
  {f'<li>Por plataforma: {plats}</li>' if plats else ''}
</ul>"""

    msg = MIMEMultipart()
    msg["Subject"] = asunto
    msg["From"]    = smtp.remitente or smtp.usuario
    msg["To"]      = ", ".join(destinatarios)
    msg.attach(MIMEText(f"""<p>Estimados,</p>
<p>Se adjuntan las estadísticas de <strong>{evento}</strong> del día <strong>{fecha_str}</strong>.</p>
{rf_block}{rs_block}
<p>Ver detalles en el archivo PDF adjunto.</p>
<hr><p style="color:#999;font-size:12px">QMS – Federación Mexicana de Radioexperimentadores A.C.</p>
""", "html"))
    att = MIMEApplication(pdf, _subtype="pdf")
    att.add_header("Content-Disposition", "attachment", filename=fname)
    msg.attach(att)

    ctx = ssl.create_default_context()
    if smtp.port == 465:
        with smtplib.SMTP_SSL(smtp.host, smtp.port, context=ctx, timeout=15) as s:
            s.login(smtp.usuario, smtp.password)
            s.sendmail(msg["From"], destinatarios, msg.as_string())
    else:
        with smtplib.SMTP(smtp.host, smtp.port, timeout=15) as s:
            s.ehlo()
            if smtp.ssl:
                s.starttls(); s.ehlo()
            s.login(smtp.usuario, smtp.password)
            s.sendmail(msg["From"], destinatarios, msg.as_string())

    return {"ok": True, "enviado_a": destinatarios, "archivo": fname}


def auto_send(db: Session, p: models.ReportePlantilla, fi: datetime, ff: datetime, user_config=None):
    """Llamado por el scheduler. Genera PDF y envía usando la config del usuario."""
    tipo = p.tipo or 'rf'
    tipo = p.tipo or 'rf'
    data: dict = {}
    if tipo in ('rf', 'ambos'):
        data['rf'] = _gather_rf(db, list(p.eventos_rf_ids or []), fi, ff)
    if tipo in ('rs', 'ambos'):
        data['rs'] = _gather_rs(db, list(p.eventos_rs_ids or []), fi, ff)
    pdf = _build_pdf(p, data, fi, ff)
    destinatarios = (user_config.destinatarios if user_config else None) or p.destinatarios or []
    _send_email_report(p, pdf, data, fi, ff, db, destinatarios_override=destinatarios)


def _resolver_fechas_ultimo_evento(db: Session, p: models.ReportePlantilla, before_date=None):
    """Resuelve fi/ff del último clúster de eventos para una plantilla.
    before_date: si se indica, busca el último evento ANTERIOR a esa fecha (para scheduler)."""
    from datetime import date as dt_date
    tipo = p.tipo or 'rf'
    fi_str = ff_str = None

    if tipo in ('rf', 'ambos'):
        rf = _cluster_rf(db, list(p.eventos_rf_ids or []), before_date=before_date)
        if rf and rf['fi']:
            fi_str = rf['fi']
            ff_str = rf['ff']

    if tipo in ('rs', 'ambos'):
        rs = _cluster_rs(db, list(p.eventos_rs_ids or []), before_date=before_date)
        if rs and rs['fi']:
            fi_str = min(fi_str, rs['fi']) if fi_str else rs['fi']
            ff_str = max(ff_str, rs['ff']) if ff_str else rs['ff']

    if not fi_str:
        return None, None

    fi = datetime.combine(dt_date.fromisoformat(fi_str), datetime.min.time())
    ff = datetime.combine(dt_date.fromisoformat(ff_str), datetime.max.time())
    return fi, ff


# ─── Endpoints generar / enviar ───────────────────────────────────────────────

@router.post("/generar/{pid}")
def generar_pdf(
    pid: int,
    fecha_inicio: datetime,
    fecha_fin: datetime,
    db: Session = Depends(get_db),
    _=Depends(get_current_user),
):
    p = db.get(models.ReportePlantilla, pid)
    if not p:
        raise HTTPException(404, "Plantilla no encontrada")

    tipo = p.tipo or 'rf'
    data: dict = {}
    if tipo in ('rf', 'ambos'):
        data['rf'] = _gather_rf(db, list(p.eventos_rf_ids or []), fecha_inicio, fecha_fin)
    if tipo in ('rs', 'ambos'):
        data['rs'] = _gather_rs(db, list(p.eventos_rs_ids or []), fecha_inicio, fecha_fin)

    pdf = _build_pdf(p, data, fecha_inicio, fecha_fin)

    fname  = f"qms_{p.nombre.replace(' ', '_').lower()}_{fecha_inicio.strftime('%Y%m%d')}.pdf"

    return StreamingResponse(
        io.BytesIO(pdf),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.post("/enviar/{pid}")
def enviar_pdf(
    pid: int,
    fecha_inicio: datetime,
    fecha_fin: datetime,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    p = db.get(models.ReportePlantilla, pid)
    if not p:
        raise HTTPException(404, "Plantilla no encontrada")

    uc = db.query(models.ReportePlantillaUserConfig).filter_by(
        plantilla_id=pid, usuario_id=current_user.id
    ).first()
    destinatarios = (uc.destinatarios if uc else None) or p.destinatarios or []
    if not destinatarios:
        raise HTTPException(400, "No tienes destinatarios configurados para este reporte")

    tipo = p.tipo or 'rf'
    data: dict = {}
    if tipo in ('rf', 'ambos'):
        data['rf'] = _gather_rf(db, list(p.eventos_rf_ids or []), fecha_inicio, fecha_fin)
    if tipo in ('rs', 'ambos'):
        data['rs'] = _gather_rs(db, list(p.eventos_rs_ids or []), fecha_inicio, fecha_fin)
    pdf = _build_pdf(p, data, fecha_inicio, fecha_fin)
    try:
        return _send_email_report(p, pdf, data, fecha_inicio, fecha_fin, db, destinatarios_override=destinatarios)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"Error al enviar correo: {e}")


@router.post("/generar-word/{pid}")
def generar_word(
    pid: int,
    fecha_inicio: datetime,
    fecha_fin: datetime,
    db: Session = Depends(get_db),
    _=Depends(get_current_user),
):
    p = db.get(models.ReportePlantilla, pid)
    if not p:
        raise HTTPException(404, "Plantilla no encontrada")
    tipo = p.tipo or 'rf'
    data: dict = {}
    if tipo in ('rf', 'ambos'):
        data['rf'] = _gather_rf(db, list(p.eventos_rf_ids or []), fecha_inicio, fecha_fin)
    if tipo in ('rs', 'ambos'):
        data['rs'] = _gather_rs(db, list(p.eventos_rs_ids or []), fecha_inicio, fecha_fin)
    docx_bytes = _build_word(p, data, fecha_inicio, fecha_fin)
    fname  = f"qms_{p.nombre.replace(' ', '_').lower()}_{fecha_inicio.strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )
